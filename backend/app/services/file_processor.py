"""
Serviço para extrair conteúdo de arquivos (PDF, DOC, Imagens)
"""
import os
from pathlib import Path
from typing import Dict, List
import PyPDF2
from docx import Document
from PIL import Image
import pytesseract

class FileProcessor:
    def __init__(self):
        pass
    
    def extract_text_from_file(self, file_path: str) -> str:
        """
        Extrai texto de arquivos baseado na extensão
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif extension in ['.doc', '.docx']:
                return self._extract_from_docx(file_path)
            elif extension in ['.png', '.jpg', '.jpeg', '.gif', '.webp']:
                return self._extract_from_image(file_path)
            elif extension == '.txt':
                return self._extract_from_txt(file_path)
            else:
                return f"[Arquivo {file_path.name} - tipo não suportado para extração]"
        except Exception as e:
            return f"[Erro ao extrair texto de {file_path.name}: {str(e)}]"
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extrai texto de PDF"""
        try:
            text_content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Limitar a 5 páginas para evitar textos muito longos
                max_pages = min(5, len(pdf_reader.pages))
                
                for page_num in range(max_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(f"--- Página {page_num + 1} ---")
                        text_content.append(text.strip())
                
                if len(pdf_reader.pages) > 5:
                    text_content.append(f"\n[PDF tem {len(pdf_reader.pages)} páginas - mostrando apenas as primeiras 5]")
            
            return '\n\n'.join(text_content) if text_content else "[PDF sem texto extraível]"
        
        except Exception as e:
            return f"[Erro ao processar PDF: {str(e)}]"
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """Extrai texto de arquivo Word (.docx)"""
        try:
            doc = Document(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            # Extrair texto de tabelas também
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(f"[Tabela] {cell.text.strip()}")
            
            return '\n\n'.join(paragraphs) if paragraphs else "[Documento Word vazio]"
        
        except Exception as e:
            return f"[Erro ao processar Word: {str(e)}]"
    
    def _extract_from_image(self, file_path: Path) -> str:
        """Extrai texto de imagem usando OCR"""
        try:
            # OCR pode não funcionar se tesseract não estiver instalado
            image = Image.open(file_path)
            
            # Tentar extrair texto usando OCR
            try:
                text = pytesseract.image_to_string(image, lang='por')
                if text.strip():
                    return f"[Texto extraído da imagem {file_path.name}]\n{text.strip()}"
                else:
                    return f"[Imagem {file_path.name} - sem texto detectado pelo OCR]"
            except pytesseract.TesseractNotFoundError:
                return f"[Imagem {file_path.name} - OCR não disponível (Tesseract não instalado)]"
        
        except Exception as e:
            return f"[Erro ao processar imagem: {str(e)}]"
    
    def _extract_from_txt(self, file_path: Path) -> str:
        """Extrai texto de arquivo .txt"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            return content if content.strip() else "[Arquivo de texto vazio]"
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    content = file.read()
                return content if content.strip() else "[Arquivo de texto vazio]"
            except Exception as e:
                return f"[Erro ao ler arquivo de texto: {str(e)}]"
        except Exception as e:
            return f"[Erro ao processar arquivo de texto: {str(e)}]"
    
    def process_multiple_files(self, file_paths: List[str]) -> str:
        """
        Processa múltiplos arquivos e retorna texto combinado
        """
        if not file_paths:
            return ""
        
        combined_content = []
        
        for file_path in file_paths:
            if os.path.exists(file_path):
                filename = Path(file_path).name
                content = self.extract_text_from_file(file_path)
                
                combined_content.append(f"=== ARQUIVO: {filename} ===")
                combined_content.append(content)
                combined_content.append("")  # Linha em branco
        
        return '\n'.join(combined_content)